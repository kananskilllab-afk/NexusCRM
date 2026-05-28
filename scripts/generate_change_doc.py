"""
Generates the Change Request & Implementation Document for Kanan Travel CRM (NexusCRM).
Combines the current application structure with the change requests raised in changes.docx.
"""

from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\Kanan\OneDrive\Desktop\Projects\tcrm\NexusCRM\Change_Request_Document.docx"

PRIMARY = RGBColor(0x1F, 0x3A, 0x68)   # navy
ACCENT  = RGBColor(0x2E, 0x86, 0xC1)   # blue
MUTED   = RGBColor(0x55, 0x55, 0x55)


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def style_run(run, *, size=11, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    style_run(r, size=18, bold=True, color=PRIMARY)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, size=14, bold=True, color=ACCENT)
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, size=12, bold=True, color=PRIMARY)
    return p


def para(doc, text, *, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, italic=italic, color=color)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.runs[0] if p.runs else p.add_run('')
    r.text = text
    style_run(r, size=11)
    return p


def make_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        style_run(r, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(hdr[i], '1F3A68')
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            style_run(r, size=10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
doc = Document()

# margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ---------- COVER ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Kanan Travel CRM (NexusCRM)")
style_run(r, size=26, bold=True, color=PRIMARY)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("Change Request & Implementation Document")
style_run(r, size=16, bold=True, color=ACCENT)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(f"Document Date: {date.today().strftime('%d %B %Y')}\n"
                 f"Prepared by: Development Team\nFor Approval / Sign-off by: Client Stakeholder")
style_run(r, size=11, italic=True, color=MUTED)

doc.add_paragraph()
para(doc,
     "This document captures (a) the current structure and capabilities of the "
     "Kanan Travel CRM application as it exists in the codebase today and "
     "(b) the consolidated list of change requests raised by the client. "
     "Each change item is tagged with the affected module, a proposed "
     "implementation approach, estimated effort, and a sign-off cell so the "
     "approver can confirm scope before development begins.",
     italic=True, color=MUTED)

doc.add_page_break()

# ---------- 1. PURPOSE ----------
h1(doc, "1. Purpose of this Document")
para(doc,
     "The client has provided a list of change requests for the existing "
     "Kanan Travel CRM. Before any development work is undertaken, this "
     "document is being shared to:")
bullet(doc, "Confirm the current scope of the application.")
bullet(doc, "Lock-in the exact list of changes that will be implemented in this revision.")
bullet(doc, "Capture sign-off from the requesting stakeholder so that scope, "
            "priority and acceptance criteria are agreed in writing.")
bullet(doc, "Serve as the reference document during UAT (User Acceptance Testing).")

# ---------- 2. CURRENT APPLICATION OVERVIEW ----------
h1(doc, "2. Current Application Overview")

h2(doc, "2.1 Technology Stack")
make_table(doc,
    ["Layer", "Technology"],
    [
        ["Frontend", "React 19, React Router 7, React Icons, ApexCharts, react-beautiful-dnd, @react-google-maps/api"],
        ["Backend", "Node.js + Express 4, JWT auth, Multer (uploads), Morgan, Mongoose"],
        ["Database", "MongoDB (via Mongoose ODM)"],
        ["Auth", "JWT-based session, role hierarchy enforced client- and server-side"],
        ["Deployment", "Client (React build) + Express API server"],
    ],
    col_widths=[4, 13])

h2(doc, "2.2 High-level Modules (Frontend Routes)")
make_table(doc,
    ["Module", "Route(s)", "Purpose"],
    [
        ["Dashboard",        "/, /dashboard",                 "KPI overview"],
        ["Leads",            "/leads, /leads/:id",            "Lead capture, lifecycle, follow-ups"],
        ["Pipeline",         "/pipeline",                     "Kanban-style lead pipeline"],
        ["Bookings",         "/bookings, /bookings/:id",      "Booking management"],
        ["Itinerary Builder","/itinerary/:id/builder",        "Trip itinerary creation"],
        ["Public Itinerary", "/it/:token",                    "Customer-facing itinerary share"],
        ["Customers / Contacts","/customers, /contacts/:id",  "Customer master & contact detail"],
        ["Finance",          "/finance/quotes, /invoices, /commissions", "Quotes, invoices, commissions"],
        ["Communications",   "/comms, /comms/sequences",      "Inbox + email sequence builder"],
        ["Documents",        "/documents",                    "Document vault"],
        ["Emails",           "/emails",                       "Email manager"],
        ["Supplier Contracts","/supplier-contracts",          "Contract repository"],
        ["Suppliers",        "/suppliers",                    "Supplier master"],
        ["Manage Catalog",   "/manage/hotels, /manage/packages","Hotel & holiday package catalog"],
        ["Reports",          "/reports",                      "Reporting & analytics"],
        ["Scheduler",        "/scheduler",                    "Calendar/scheduler"],
        ["Users",            "/users",                        "User & role administration"],
        ["Settings",         "/settings",                     "System settings"],
    ],
    col_widths=[3.5, 5.5, 8])

h2(doc, "2.3 Backend API Surface")
para(doc, "Express routers currently exposed under /api:")
bullet(doc, "auth, users  -  authentication, user CRUD, role hierarchy")
bullet(doc, "leads, customers, suppliers  -  primary CRM entities")
bullet(doc, "tasks, activity_feed, analytics  -  workflow & reporting")
bullet(doc, "travel_services, loyalty  -  travel offerings & loyalty points")
bullet(doc, "voyage/* (auth, bookings, contracts, documents, emails, finance, "
            "itinerary, passengers, pipeline, segments)  -  the next-generation "
            "'Voyage' module set")

h2(doc, "2.4 Core Data Models (MongoDB / Mongoose)")
make_table(doc,
    ["Model", "Purpose"],
    [
        ["Lead",              "Enquiry record - personal details, destination, enquiry type, traveller counts, travel dates"],
        ["Customer",          "Converted/master customer record"],
        ["Supplier / SupplierRate / AssignedSupplier", "Vendor master, rate cards, allocation"],
        ["BillingItem / Payment", "Line items billed to a lead/booking and payments received"],
        ["FollowUp / Task / Activity", "Scheduled follow-ups, tasks, and activity audit feed"],
        ["Communication / Notification", "Email/SMS/WhatsApp logs & in-app notifications"],
        ["CRMUser",           "Internal users with role-based access"],
        ["LoyaltyPoints",     "Customer loyalty balance"],
        ["AuditLog",          "Change-tracking ledger"],
        ["SavedFilter, ApiKey", "Saved table views and API access tokens"],
    ],
    col_widths=[5, 12])

h2(doc, "2.5 Current 'Add Lead' Form Capability (baseline)")
bullet(doc, "Enquiry types supported: Flight, Hotel, Visa, Package")
bullet(doc, "Captures: name, mobile, email, destination (free-text), source, "
            "adults/children, priority, travel start / travel end (free date), tags, GDPR consent")
bullet(doc, "Hotel sub-form: star rating + meal plan only (no room types, no infant/child age)")
bullet(doc, "Flight sub-form: origin + class")
bullet(doc, "No UTM / campaign capture, no qualified/unqualified flag, "
            "no credit-card details, no numeric unique code")

doc.add_page_break()

# ---------- 3. CHANGE REQUESTS ----------
h1(doc, "3. Consolidated Change Requests")
para(doc,
     "All items listed below are taken verbatim from the change list supplied "
     "by the client (changes.docx). They have been grouped, given an ID, and "
     "annotated with the affected module and proposed implementation approach. "
     "The 'Approval' column is to be initialled by the requester at sign-off.",
     italic=True, color=MUTED)

# ---- 3.1 General changes ----
h2(doc, "3.1 General / Cross-cutting Changes")
make_table(doc,
    ["ID", "Change Request", "Affected Module(s)", "Proposed Implementation", "Approval"],
    [
        ["G-01",
         "Add +/-3 day flexibility on travel start & travel end dates when creating a Flight lead.",
         "AddLeadModal, Lead model",
         "Replace single date inputs with a date + flex-window selector; persist flex_days on the lead's enquiry_data.flight payload.",
         ""],
        ["G-02",
         "Paid-campaign integration on ALL forms — capture UTM parameters.",
         "AddLeadModal, public web forms, Lead schema",
         "Read utm_source / utm_medium / utm_campaign / utm_term / utm_content from the URL and POST body; persist on the Lead record; expose in filters & reports.",
         ""],
        ["G-03",
         "Build a proper Dashboard with real, working KPIs and drill-downs.",
         "Dashboard.js, analytics API",
         "Define KPI tiles (leads, conversion, revenue, follow-ups due, RR distribution), tie each to an /api/analytics endpoint, and add drill-down navigation to filtered lists.",
         ""],
        ["G-04",
         "Destination field must be a dropdown (currently free text).",
         "AddLeadModal, LeadDetail, filters",
         "Introduce a Destinations master (City + Country); replace free-text inputs with a searchable single-select.",
         ""],
        ["G-05",
         "Collect all finance documents (rules, T&Cs, receipt formats, etc.) and integrate.",
         "Documents Vault, Finance",
         "Client to provide source files. We will template Invoice/Receipt, link T&Cs to quotes, and store master docs in the Document Vault.",
         ""],
        ["G-06",
         "Enable cookies / cookie consent.",
         "Frontend shell",
         "Add a cookie-consent banner, persist preference, integrate analytics only after consent.",
         ""],
        ["G-07",
         "Add Credit Card details capture.",
         "Customer / Payment",
         "Capture last-4, network, expiry & cardholder name only (no PAN/CVV stored). Mark fields PCI-sensitive; encrypt at rest.",
         ""],
        ["G-08",
         "Make the unique Lead Code numeric.",
         "Lead model, Lead list",
         "Replace UUID-style id with an auto-incrementing numeric counter (e.g., LD-100001). Keep internal _id for relations.",
         ""],
        ["G-09",
         "Add a 'Qualified / Unqualified' column on the lead list.",
         "Lead schema, LeadList table",
         "Add qualification: 'Qualified' | 'Unqualified' | 'Pending' to the schema; expose as a column and filter."
         , ""],
        ["G-10",
         "If Qualified ⇒ automatically create a Follow-up.",
         "Leads workflow",
         "On qualification change to 'Qualified', auto-create a FollowUp with a default SLA (configurable in Settings).",
         ""],
        ["G-11",
         "Text/Comms compliance — capture customer GST details.",
         "Customer schema, billing",
         "Add gstin, legal_name and place_of_supply to the customer record; surface on invoices.",
         ""],
        ["G-12",
         "Get template approvals (WhatsApp/SMS).",
         "Comms / Templates",
         "Add a template registry with status (Draft / Submitted / Approved / Rejected) and provider template-id; block sends on non-approved templates.",
         ""],
        ["G-13",
         "All tabular data must be filterable.",
         "Every table view",
         "Standardise on the existing SavedFilter model; add column filters (text/select/date-range) to Leads, Customers, Suppliers, Bookings, Invoices, Quotes, Commissions.",
         ""],
        ["G-14",
         "Approval process for discounts above a defined threshold.",
         "Quotes / Billing",
         "Add a discount_approval workflow: if discount > threshold (per setting), the quote moves to 'Pending Approval' and requires a Manager+ role to release.",
         ""],
        ["G-15",
         "Remove Margin field from visible UI.",
         "Quotes / Billing UI",
         "Hide margin column from non-Admin roles; retain in DB for reporting only.",
         ""],
        ["G-16",
         "Review billing module end-to-end.",
         "Finance (Quotes, Invoices, Commissions, BillingItem, Payment)",
         "Walkthrough session with Vickey Sir to confirm invoice format, GST, taxes, advance vs balance, refunds and receipts; produce a small follow-up addendum.",
         ""],
        ["G-17",
         "Meet Vickey Sir for billing data — hotels, flights, etc.",
         "Project / Stakeholder",
         "Schedule a working session; capture sample invoices for each line of business; feed into G-05 and G-16.",
         ""],
        ["G-18",
         "Contacts menu should appear below Leads in the sidebar.",
         "Sidebar",
         "Re-order the menu items array so that Contacts renders directly below Leads.",
         ""],
        ["G-19",
         "Lead-score logic should be even / balanced.",
         "Lead scoring",
         "Re-weight scoring so each criterion contributes equal points; document the formula in Settings.",
         ""],
        ["G-20",
         "Can the Lead Score be trained by Gemini (AI)?",
         "Lead scoring",
         "Spike: integrate Google Gemini API to suggest a score based on lead attributes. Treat as Phase-2 (feasibility + cost confirmation required).",
         ""],
        ["G-21",
         "Round-robin assignment of leads.",
         "Leads workflow",
         "Add a round-robin assignment service that rotates new leads across eligible agents (respecting capacity & working hours).",
         ""],
        ["G-22",
         "Every change MUST be tested before submission (QA gate).",
         "Process",
         "Adopt a written test plan per change item; no item is signed off until UAT is recorded against it.",
         ""],
    ],
    col_widths=[1.2, 5, 3, 6.5, 1.3])

# ---- 3.2 Leads - Hotel ----
h2(doc, "3.2 Leads → Hotel Sub-form")
make_table(doc,
    ["ID", "Change Request", "Proposed Implementation", "Approval"],
    [
        ["H-01",
         "Add 'Number of rooms' next to star rating, with room type options: Double sharing, Twin sharing, Quad sharing, Oblique/Family room, Suite.",
         "Add a 'rooms' array on enquiry_data.hotel = [{ type, count }]; render a small repeater UI beside the star rating.",
         ""],
        ["H-02",
         "Capture age for each child in the Children section.",
         "Replace simple child count with a list of child ages; persist in enquiry_data.hotel.children_ages.",
         ""],
        ["H-03",
         "Rename 'Travel Date / Return Date' to 'Check-in / Check-out' for hotel enquiries.",
         "Conditional labels: when activeType === 'Hotel', show Check-in / Check-out; otherwise keep Travel Start / Travel End.",
         ""],
        ["H-04",
         "Set up Kanan Travel social media integration in the Hotel flow.",
         "Clarify with client: scope = lead capture from social channels (FB/IG/WhatsApp) into the hotel-lead pipeline. Track as a discovery item.",
         ""],
        ["H-05",
         "Fix: 'New Holiday Packages are not adding' (bug).",
         "Reproduce on /manage/packages — verify form validation, payload, and POST /api/travel_services; patch the failing endpoint/validator.",
         ""],
    ],
    col_widths=[1.2, 5, 9.5, 1.3])

# ---- 3.3 Leads - Visa ----
h2(doc, "3.3 Leads → Visa Sub-form")
make_table(doc,
    ["ID", "Change Request", "Proposed Implementation", "Approval"],
    [
        ["V-01",
         "Add a 'Type of Visa' column / field.",
         "Add visa_type (Tourist, Business, Student, Work, Transit, Medical, Other) to enquiry_data.visa; show as a column on the lead list and as a filter.",
         ""],
    ],
    col_widths=[1.2, 5, 9.5, 1.3])

# ---- 3.4 Leads - Packages ----
h2(doc, "3.4 Leads → Packages Sub-form")
make_table(doc,
    ["ID", "Change Request", "Proposed Implementation", "Approval"],
    [
        ["P-01",
         "Destination must be a multi-select dropdown on Package enquiries.",
         "Replace the single destination input with a multi-select bound to the Destinations master (see G-04); persist as an array.",
         ""],
    ],
    col_widths=[1.2, 5, 9.5, 1.3])

doc.add_page_break()

# ---------- 4. OUT OF SCOPE / OPEN ITEMS ----------
h1(doc, "4. Items Needing Client Clarification")
para(doc, "The following items are accepted in principle but need additional input from the requester before sizing:")
bullet(doc, "G-05 / G-17: Source finance documents (rules, T&Cs, receipt formats) must be shared by the client.")
bullet(doc, "G-12: WhatsApp/SMS template texts and provider (Gupshup, Twilio, MSG91, etc.) to be confirmed.")
bullet(doc, "G-14: Discount threshold value(s) and approver role to be specified.")
bullet(doc, "G-20: Gemini-based scoring is a feasibility spike; final scope and cost to be agreed after the spike.")
bullet(doc, "H-04: 'Travel with Kanan' social media set-up - scope of integration to be confirmed.")

# ---------- 5. DELIVERY ASSUMPTIONS ----------
h1(doc, "5. Delivery Assumptions")
bullet(doc, "Each change item will be tested before being marked complete (per G-22).")
bullet(doc, "Sign-off on this document freezes the scope of this revision. New items raised after sign-off will be tracked as a separate change request.")
bullet(doc, "Phase-2 items (e.g., G-20) will be quoted separately once scope is confirmed.")
bullet(doc, "No PAN/CVV will be stored under G-07 — only PCI-safe card metadata.")

# ---------- 6. APPROVAL ----------
h1(doc, "6. Approval & Sign-off")
para(doc,
     "By signing below, the requester confirms that the change list captured "
     "in Section 3 is complete and correctly reflects their requirements, and "
     "authorises the development team to begin implementation accordingly.",
     italic=True)

doc.add_paragraph()
make_table(doc,
    ["Role", "Name", "Signature", "Date"],
    [
        ["Requested by (Client)", "", "", ""],
        ["Reviewed by (Project Manager)", "", "", ""],
        ["Accepted by (Development Lead)", "", "", ""],
    ],
    col_widths=[4.5, 4, 4, 3.5])

doc.add_paragraph()
para(doc,
     "Notes / additional remarks from the requester:",
     bold=True)
for _ in range(6):
    doc.add_paragraph("_" * 95)

doc.save(OUTPUT)
print(f"Document generated: {OUTPUT}")
