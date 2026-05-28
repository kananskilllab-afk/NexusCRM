import os
import sys
import traceback
from docx import Document
from docx.shared import Inches, Pt

# Path setup
md_file = r"C:\Users\mahim\.gemini\antigravity\brain\b44dcbb2-1c17-4099-9616-60811f5e6fdb\Requirement_Document.md.resolved"
output_docx = r"c:\Users\mahim\.gemini\antigravity\scratch\TCRM\TravoCRM_Requirement_Document.docx"
image_path = r"C:\Users\mahim\.gemini\antigravity\brain\b44dcbb2-1c17-4099-9616-60811f5e6fdb\mermaid_diagram_complete_1776671255059.png"

def create_docx():
    try:
        doc = Document()
        
        # Heading
        doc.add_heading("TravoCRM: Comprehensive System Requirements & Architecture Document", 0)

        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()

        sections = content.split('\n\n')
        
        for section in sections:
            lines = section.strip().split('\n')
            if not lines:
                continue
                
            first_line = lines[0].strip()
            
            # Handle Headings
            if first_line.startswith('### '):
                doc.add_heading(first_line.replace('### ', ''), level=3)
                lines = lines[1:]
            elif first_line.startswith('## '):
                doc.add_heading(first_line.replace('## ', ''), level=2)
                lines = lines[1:]
            elif first_line.startswith('# '):
                doc.add_heading(first_line.replace('# ', ''), level=1)
                lines = lines[1:]

            # Handle Table
            if '|' in first_line:
                rows_data = []
                for line in lines:
                    if '|' in line and not any(c in line for c in [':---', '---']):
                        cells = [c.strip().replace('**', '') for c in line.split('|') if c.strip()]
                        if cells:
                            rows_data.append(cells)
                
                if rows_data:
                    num_of_cols = max(len(row) for row in rows_data)
                    table = doc.add_table(rows=len(rows_data), cols=num_of_cols)
                    table.style = 'Light Shading Accent 1' # Use a safer style
                    for i, row_data in enumerate(rows_data):
                        for j, cell_text in enumerate(row_data):
                            if j < num_of_cols:
                                table.cell(i, j).text = cell_text
                continue

            # Handle Mermaid
            if '```mermaid' in first_line or any('```mermaid' in l for l in lines):
                doc.add_paragraph("Operational Workflow Diagram:")
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(6))
                else:
                    doc.add_paragraph("[Diagram Image Missing]")
                continue

            # Handle Lists and Paragraphs
            for line in lines:
                l = line.strip()
                if l.startswith('* ') or l.startswith('- '):
                    doc.add_paragraph(l[2:], style='List Bullet')
                elif l.startswith('1. ') or l.startswith('2. '):
                    doc.add_paragraph(l[3:], style='List Number')
                elif l:
                    doc.add_paragraph(l.replace('**', ''))

        doc.save(output_docx)
        print(f"Success: {output_docx}")
    except Exception as e:
        print(f"Error: {str(e)}")
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    create_docx()
